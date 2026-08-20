from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
import os

OUT = r"D:\PROJECT\PicturePuzzle_Framework\output\docx\PicturePuzzle_Framework_Developer_Guide.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "5B6573"
LIGHT = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
GREEN = "E8F5E9"
AMBER = "FFF4D6"
RED = "FCE8E6"
WHITE = "FFFFFF"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_border(cell, color="B8C2CC", sz="8", val="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), sz)
        element.set(qn("w:color"), color)

def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_cell_height(cell, height_inches):
    tcPr = cell._tc.get_or_add_tcPr()
    trPr = cell._tc.getparent().get_or_add_trPr()
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(int(height_inches * 1440)))
    trHeight.set(qn("w:hRule"), "atLeast")

def keep_with_next(p):
    p.paragraph_format.keep_with_next = True

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def style_run(run, bold=False, color=None, size=None, italic=False, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)

def add_field_update(settings):
    settings = settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.line_spacing = 1.18
normal.paragraph_format.space_after = Pt(5)

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 8),
    ("Heading 2", 13, BLUE, 13, 6),
    ("Heading 3", 11.5, DARK_BLUE, 9, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "CodeBlock" not in [s.name for s in styles]:
    code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["CodeBlock"]
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor.from_string("263238")
code_style.paragraph_format.left_indent = Inches(0.08)
code_style.paragraph_format.right_indent = Inches(0.08)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.0

if "CaptionCustom" not in [s.name for s in styles]:
    cap = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
else:
    cap = styles["CaptionCustom"]
cap.font.name = "Calibri"
cap.font.size = Pt(8.5)
cap.font.italic = True
cap.font.color.rgb = RGBColor.from_string(MUTED)
cap.paragraph_format.space_after = Pt(7)

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.28)
    st.paragraph_format.first_line_indent = Inches(-0.14)
    st.paragraph_format.space_after = Pt(2)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("PICTURE PUZZLE FRAMEWORK  /  INTERNAL DEVELOPER GUIDE")
style_run(r, bold=True, color=MUTED, size=7.5)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Tài liệu kỹ thuật nội bộ  •  Unity 6000.3.12f1  •  Trang ")
style_run(r, color=MUTED, size=8)
add_page_number(footer)
add_field_update(doc.settings)

def para(text="", bold_prefix=None, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        style_run(a, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        style_run(b)
    else:
        p.add_run(text)
    return p

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.28 + 0.22 * level)
    return p

def number(text):
    return doc.add_paragraph(text, style="List Number")

def code(text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    set_cell_border(cell, color="CBD5E1", sz="6")
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    p = cell.paragraphs[0]
    p.style = "CodeBlock"
    p.paragraph_format.keep_together = True
    p.add_run(text)
    return table

def callout(title, text, fill=LIGHT_BLUE, border=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, sz="8")
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    a = p.add_run(title)
    style_run(a, bold=True, color=border, size=10)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    return table

def image_placeholder(num, title, purpose, hint, height=1.65):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    set_cell_border(cell, color="9AA8B5", sz="10", val="dashed")
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    set_cell_height(cell, height)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[CHÈN ẢNH %02d]\n%s" % (num, title))
    style_run(r, bold=True, color=MUTED, size=10)
    p2 = cell.add_paragraph("Mục đích: " + purpose)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p2.runs[0], color=MUTED, size=8.5)
    p3 = cell.add_paragraph("Gợi ý: " + hint)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p3.runs[0], italic=True, color=MUTED, size=8)
    cap = doc.add_paragraph("Hình %02d — %s" % (num, title), style="CaptionCustom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def table(headers, rows, widths=None, header_fill=BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths is None:
        widths = [6.55 / len(headers)] * len(headers)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color=header_fill, sz="6")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, bold=True, color=WHITE, size=9)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_border(cells[i], color="CBD5E1", sz="5")
            set_cell_margins(cells[i], top=90, start=110, bottom=90, end=110)
            set_cell_shading(cells[i], "FFFFFF" if ridx % 2 == 0 else "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            style_run(r, size=8.6)
    return t

def page_break():
    doc.add_page_break()

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(34)
r = p.add_run("PICTURE PUZZLE")
style_run(r, bold=True, color=BLUE, size=12)
p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after = Pt(3)
r = p2.add_run("Framework Developer Guide")
style_run(r, bold=True, color=DARK_BLUE, size=27)
p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(20)
r = p3.add_run("Tài liệu vận hành, kiến trúc và mở rộng SDK cho đội phát triển")
style_run(r, color=MUTED, size=14)
callout("PHẠM VI TÀI LIỆU", "Tài liệu này mô tả source và scene hiện tại của PicturePuzzle_Framework. Các đoạn ghi “UI-only”, “local fallback” hoặc “optional provider” là ranh giới kỹ thuật: không được suy diễn thành gameplay puzzle runtime đã có.", LIGHT_BLUE, BLUE)
para("")
table(["Thông tin", "Giá trị"], [
    ("Unity", "6000.3.12f1"),
    ("Build scenes", "Init → Menu → Game"),
    ("Đối tượng đọc", "Game/UI/SDK developers, technical artists, QA"),
    ("Phiên bản tài liệu", "Developer handbook / Word draft"),
], [1.55, 5.0])
para("")
para("Cách dùng các khung ảnh", bold_prefix="Cách dùng các khung ảnh")
para("Mỗi khung [CHÈN ẢNH xx] là vùng cố ý để team dán screenshot từ Unity Inspector, Hierarchy, Game view hoặc Console. Giữ caption và số hình để tài liệu PDF sau này có thể tham chiếu ổn định.")
callout("QUY ƯỚC ĐỌC NHANH", "Tên class, prefab, scene và đường dẫn asset được giữ nguyên tiếng Anh. Tên trạng thái “Core chính thức”, “UI/local fallback”, “Optional provider”, “UI-only” và “Chưa có gameplay runtime” là nhãn dùng xuyên suốt handbook.", AMBER, "B7791F")
page_break()

# TOC and status
heading("Mục lục và bản đồ đọc", 1)
para("Đọc theo thứ tự nếu bạn mới vào project; đọc phần 5 và 6 khi cần tích hợp module hoặc puzzle runtime.")
table(["Phần", "Nội dung", "Khi cần đọc"], [
    ("1", "Luồng khởi động: Init, GameLoading, Initializer", "Debug boot, thêm service nền"),
    ("2", "Project Init Settings và thứ tự module", "Cấu hình module, tạo InitModule"),
    ("3", "Menu scene và hệ thống popup", "Sửa level map, popup, bottom navigation"),
    ("4", "Game scene và UI-only shell", "Bind HUD, complete/game-over, Power Up"),
    ("5", "Tổ chức module và cách dùng", "Tích hợp Save, Audio, Currency, Quest…"),
    ("6", "Contract puzzle và mở rộng", "Gắn board/level/session của game team"),
    ("7", "Vận hành, build và smoke test", "QA, release, xử lý Console"),
], [0.55, 3.7, 2.3])
heading("Ma trận trạng thái framework", 2)
table(["Nhãn", "Ý nghĩa trong project này", "Ví dụ"], [
    ("Core chính thức", "Runtime nền tảng, được khởi tạo từ NebulaSoft Core.", "Initializer, Save, Tween, Audio, UI"),
    ("UI/local fallback", "UI và service local để scene chạy mà không cần backend thật.", "Local leaderboard, Profile, Quest UI"),
    ("Optional provider", "Adapter có thể bật/tắt; không được buộc Core phụ thuộc.", "Ads, Firebase, IAP, analytics provider"),
    ("UI-only", "Có prefab/visual binding nhưng không xử lý gameplay/data thật.", "Power Up bar, Game HUD, LevelData shell"),
    ("Chưa có gameplay runtime", "Không có engine board/tile/rule trong SDK hiện tại.", "LevelController, Board, Tile, puzzle rules"),
], [1.25, 3.25, 2.05])
callout("KẾT LUẬN KIẾN TRÚC", "SDK hiện đủ để làm framework shell cho một dòng puzzle: boot, navigation, popup, HUD, reward/currency và extension point. Game team vẫn phải cung cấp level authoring, board, input, session, win/fail và luật puzzle.", RED, "B42318")
page_break()

# Part 1
heading("1. Luồng khởi động: từ Init đến scene đầu tiên", 1)
para("Build Settings hiện có ba scene bật theo thứ tự: Assets/Project Files/Game/Scenes/Init.unity, Menu.unity và Game.unity. Khi chạy build hoặc Play Mode từ Init, hệ thống không nhảy thẳng vào Menu; nó đi qua GameLoading để dựng runtime nền tảng và kiểm tra kết nối.")
heading("1.1. Init scene chứa gì?", 2)
para("Init.unity là scene bootstrap. Thành phần quan trọng là GameLoading và các reference serialized trỏ tới Initializer, loading graphics, optional NetworkConnection probe và scene build index. Init không phải gameplay scene; trách nhiệm của nó là tạo một runtime sống qua scene bằng DontDestroyOnLoad rồi chuyển sang scene kế tiếp.")
bullet("GameLoading.Awake() lưu static instance, gọi loadingGraphics.Init(this) và bắt đầu BootstrapCoroutine().")
bullet("BootstrapCoroutine() đợi một frame và EndOfFrame để Unity hoàn tất Awake/scene binding trước khi init.")
bullet("Initializer.Init() chạy một lần, bind EventSystem/input module, SystemMessage, Overlay, AnalyticsModules và đánh dấu object DontDestroyOnLoad.")
bullet("Initializer.InitModules() đọc ProjectInitSettings và gọi CreateComponent() theo array modules.")
bullet("Initializer.InitSDKs() gọi SDKInitializer.Init() cho các SDK/provider tùy chọn.")
image_placeholder(1, "Init scene / GameLoading Inspector", "Cho thấy object GameLoading, Initializer reference và loading graphics.", "Chụp Hierarchy + Inspector của Init.unity; khoanh vùng scene build index và các reference.", 1.6)
heading("1.2. Trình tự runtime chính xác", 2)
number("Unity load Init.unity; GameLoading.Awake() chạy trước scene gameplay.")
number("BootstrapCoroutine() đợi một frame + EndOfFrame.")
number("initializer.Init() chạy guard một lần. Nếu đã init, lần gọi sau không làm lại static services.")
number("ConnectionCheckCoroutine() đặt loading message “Checking connection..”. Nếu checkNetworkConnection bật, probe ServiceProbeUrl; lỗi thì hiện “Connection error” và dừng ở loading scene.")
number("Nếu kết nối đạt, tiến độ về 0.1 với message “Loading..”; gọi InitModules(), sau đó InitSDKs().")
number("Các loading task tĩnh được activate lần lượt; GameLoading đợi IsFinished của từng task.")
number("Sau vài frame ổn định, tính scene kế tiếp theo Build Settings hoặc LoadingSceneBuildIndex.")
number("LoadSceneAsync() chạy với tiến độ hiển thị trong khoảng 0.2–0.9; giữ thời gian loading tối thiểu hai giây.")
number("Đặt message “Completed”, chờ MarkAsReadyToHide() nếu manual control bật, gọi loadingGraphics.OnLoadingFinished(), rồi destroy GameLoading.")
code("""// Điểm vào khái niệm (không gọi trực tiếp từ gameplay)
GameLoading.AddTask(myTask);
GameLoading.SetLoadingMessage("Loading puzzle content");
GameLoading.MarkAsReadyToHide();""")
heading("1.3. Bổ sung thứ gì ở Init và bổ sung ở đâu?", 2)
table(["Nhu cầu", "Nơi nên thêm", "Không nên làm"], [
    ("Service nền cần có trước Menu", "Tạo InitModule trong Core hoặc project feature; thêm asset module vào Project Init Settings.", "Khởi tạo singleton trong MenuController.Start()."),
    ("SDK/provider tùy chọn", "SDKInitializer hoặc adapter riêng; giữ provider sau Core Init.", "Cho Core gọi trực tiếp Firebase/Ads SDK."),
    ("Loading task bất đồng bộ", "Đăng ký qua GameLoading.AddTask() và expose IsFinished.", "Chặn Main Thread bằng vòng lặp vô hạn."),
    ("UI scene cụ thể", "Menu/Game UIController và prefab tương ứng.", "Đưa UI scene dependency vào Init scene."),
], [1.55, 3.2, 1.8])
callout("NGUYÊN TẮC", "Init chỉ dựng nền tảng. Menu quyết định navigation; Game quyết định HUD/session shell. Nếu một service chỉ phục vụ một popup, cân nhắc lazy init tại page thay vì làm InitModule.", LIGHT_BLUE, BLUE)
image_placeholder(3, "Init loading screen / Console timeline", "Đối chiếu thứ tự message Checking connection → Loading → Completed với Console.", "Chụp Game view loading và Console cùng thời điểm Play Mode.", 1.45)

heading("1.4. Initializer, static state và an toàn khi đổi scene", 2)
para("Initializer có DefaultExecutionOrder(-999), giữ InitSettings static và được DontDestroyOnLoad. Điều này giúp Menu và Game truy cập cùng runtime, nhưng cũng làm lỗi state khó thấy nếu service tự giữ reference đến object của scene cũ.")
bullet("Static service chỉ giữ dữ liệu/runtime dùng xuyên scene; không giữ Transform hoặc UI component của Menu/Game nếu không có cơ chế clear.")
bullet("Popup phải đóng event ở OnDisable/CompleteHide một lần; không phát lại event trong OnDestroy khi object đang bị hủy.")
bullet("Khi đổi scene, unsubscribe event trước khi destroy provider hoặc page; không gọi SetActive() lên object đang destroy.")
bullet("Nếu cần reset state cho test, dùng API reset rõ ràng hoặc reload domain; không tự ý thêm DontDestroyOnLoad cho prefab UI.")

# Part 2
page_break()
heading("2. Project Init Settings: cấu hình và thứ tự chạy", 1)
para("Asset thật là Assets/Project Files/Data/Project Init Settings.asset. ProjectInitSettings chứa serialized InitModule[] modules. ProjectInitSettings.Init(initializer) lặp array và gọi CreateComponent() trên từng module khác null; GetModule<T>() dùng để truy vấn module đã gắn.")
image_placeholder(2, "Project Init Settings Inspector", "Cho thấy danh sách Modules và thứ tự cấu hình mà team đang nhìn thấy trong Unity.", "Dán screenshot Inspector do team chụp; có thể dùng ảnh clipboard đã cung cấp, nhưng giữ khung để thay ảnh có độ phân giải cao.", 2.15)
heading("2.1. Thứ tự module trong asset hiện tại", 2)
table(["#","Module / RegisterModule","Cấu hình đang gắn","Trạng thái"], [
    ("1","Save Controller / SaveInitModule","autoSaveDelay 0; cleanSaveStart 0; WebGL prefix build_20260223_173146","Core chính thức"),
    ("2","Tween / TweenInitModule","300 update; 30 fixed; 0 late; verbose off","Core chính thức"),
    ("3","Audio Controller / AudioInitModule","Audio settings; pool 4; max distance 30; spread 180","Core chính thức"),
    ("4","Currencies / CurrencyInitModule","Currencies database","Core chính thức"),
    ("5","Haptic / HapticInitModule","verbose off","Core chính thức"),
    ("6","Dev Panel / DevPanelInitModule","Dev panel settings","UI/local fallback"),
    ("7","Game Settings / GameInitModule","GameData","UI/local fallback"),
    ("8","Screen Settings / ScreenSettings","60 FPS; battery 30; sleep timeout -1","Core chính thức"),
    ("9","Lives System / LivesSystemInitModule","LivesData","UI/local fallback"),
    ("10","Quest / QuestInitModule","QuestDatabase","UI/local fallback"),
    ("11","Daily Reward / DailyRewardInitModule","DailyRewardDatabase","UI/local fallback"),
], [0.25, 1.65, 2.85, 1.8])
para("Lưu ý: Inspector có thể nhóm hoặc sắp xếp nhãn theo custom editor, nhưng runtime dùng thứ tự serialized trong asset. RegisterModuleAttribute chỉ cung cấp tên, path, core flag và order cho editor/discovery; attribute không tự tạo component nếu module chưa được add vào asset.")
heading("2.2. Mẫu thêm InitModule mới", 2)
para("Khi puzzle team cần một runtime service, tạo module nhỏ ở project feature, tạo config asset riêng, rồi add module vào Project Init Settings. Không đặt logic board trong Initializer.")
code("""[RegisterModule("Puzzle", core: false)]
public sealed class PuzzleInitModule : InitModule
{
    [SerializeField] private PuzzleConfig config;
    public override string ModuleName => "Puzzle";

    public override void CreateComponent()
    {
        PuzzleRuntime.Initialize(config);
    }
}""")
number("Tạo class kế thừa InitModule trong module assembly phù hợp.")
number("Nếu muốn hiện trong custom editor, thêm RegisterModule(\"Puzzle\", core:false).")
number("Tạo PuzzleConfig.asset và kéo vào field config.")
number("Mở Assets/Project Files/Data/Project Init Settings.asset; thêm PuzzleInitModule vào Modules.")
number("Chạy từ Init, kiểm tra log khởi tạo trước khi Menu được load.")
code("""var settings = Initializer.InitSettings;
var gameSettings = settings.GetModule<GameInitModule>();
// Chỉ đọc config/runtime; không tự tạo module lần thứ hai.""")
image_placeholder(9, "Module configuration example in Inspector", "Minh họa field config, asset reference và module mới sau khi add vào Modules.", "Chụp Inspector của InitModule asset và Project Init Settings cùng lúc.", 1.55)
image_placeholder(10, "Custom InitModule code và ProjectInitSettings asset", "Đặt cạnh ví dụ code để dev mới hiểu quan hệ giữa class, config và serialized asset.", "Dán screenshot IDE + Inspector; che thông tin nhạy cảm của provider nếu có.", 1.55)

heading("2.3. Sai lầm thường gặp khi cấu hình Init", 2)
table(["Triệu chứng", "Nguyên nhân có thể", "Cách kiểm tra / sửa"], [
    ("Module row có trong editor nhưng runtime không chạy", "Chưa add instance vào modules array hoặc asset đang dùng không phải asset trong Init.", "Chọn Init scene, kiểm tra Initializer.initSettings và array serialized."),
    ("NullReference khi init", "Config database/prefab chưa kéo vào field.", "Mở module asset, kiểm tra GUID/reference; thêm guard log rõ tên asset."),
    ("Menu chạy nhưng service reset mỗi scene", "Service được tạo trong scene controller thay vì InitModule.", "Chuyển init về ProjectInitSettings nếu state cần xuyên scene."),
    ("Provider làm Core compile fail", "Core reference trực tiếp package/SDK tùy chọn.", "Đưa vào adapter/assembly optional và gọi qua SDKInitializer."),
], [1.55, 2.35, 2.65])

# Part 3
page_break()
heading("3. Menu scene: navigation, level map và popup", 1)
para("Menu.unity là scene đầu tiên sau Init. Hierarchy hiện có Bottom Navigation Canvas, UI Main Canvas, Scripts Holder, Tutorial Overlay, Fade, Light, Directional Light và Main Camera. MenuController là orchestration layer; UIController là page registry và lifecycle manager.")
image_placeholder(4, "Menu scene hierarchy", "Xác nhận tên object top-level và vị trí của UIController/MenuController.", "Chụp Hierarchy mở rộng UI Main Canvas, Bottom Navigation Canvas và Scripts Holder.", 1.7)
heading("3.1. MenuController chạy như thế nào?", 2)
code("""private void Awake()
{
    uiController?.Init();
    uiController?.InitPages();
}

private void Start()
{
    UIController.ShowPage<UIMainMenu>();
    _ = LocalLeaderboardService.PreloadLeaderboardsAsync();
    AdsManager.EnableBanner();
    Overlay.Hide(0.3f);
}""")
bullet("Awake đăng ký toàn bộ UIPage con và chuẩn bị cache/popup.")
bullet("Start mở UIMainMenu, preload local leaderboard, bật banner qua AdsManager và ẩn Overlay.")
bullet("MenuController.LoadGame() chỉ load Game nếu SceneUtils.DoesSceneExist(GameConsts.SCENE_GAME) trả true.")
bullet("Không thêm logic tạo board/level vào MenuController; Menu chỉ phát lệnh navigation và render map.")
heading("3.2. UIController và lifecycle page", 2)
para("UIController.Init() quét các UIPage con, gọi CacheComponents() và tạo pagesLink theo Type. InitPages() cấu hình CanvasScaler, NotchSaveArea, FloatingCloud, rồi PreparePage()/Init() từng page. ShowPage<T>() bật GraphicRaycaster, chạy PlayShowAnimation và EnableCanvas; HidePage<T>() tắt raycaster và chạy PlayHideAnimation.")
table(["Giai đoạn", "Trách nhiệm", "Điểm cần giữ"], [
    ("Init", "Tìm UIPage, cache component, phát hiện duplicate type.", "Mỗi page type chỉ có một registration."),
    ("InitPages", "Safe area, CanvasScaler, currency cloud, page.Init().", "Không đọc data gameplay chưa có."),
    ("ShowPage", "Bật canvas/raycaster, gọi animation và OnPageOpened.", "Popup pause có thể set Time.timeScale=0."),
    ("HidePage", "Tắt raycaster, chạy hide animation.", "Đóng event một lần ở CompleteHide/OnDisable."),
    ("OnPageClosed", "Unregister popup, khôi phục time scale khi không còn pause popup.", "Không gọi lại từ OnDestroy."),
], [1.15, 3.55, 1.85])
image_placeholder(5, "UI Main Menu level map", "Cho thấy Levels Map, Level Point, Play Button, progress box và avatar.", "Chụp Game view Menu và Hierarchy prefab UI Main Menu.", 1.75)
heading("3.3. Các popup/page hiện có", 2)
table(["Popup / page", "Prefab / entry point", "Vai trò và trạng thái"], [
    ("UIMainMenu", "UI Main Menu.prefab", "Level map, profile/settings/no-ads buttons; UI shell."),
    ("UIProfilePopup", "Assets/Addon/UI/Prefabs/Pages/UI Profile Popup.prefab", "Profile; local UI. OnDisable chịu trách nhiệm trả bottom navigation."),
    ("Settings", "UI Menu Settings.prefab", "Screen/audio/settings UI; đọc config hiện có."),
    ("No Ads", "UI No Ads Pop Up.prefab", "Monetization offer; provider có thể optional."),
    ("Leaderboard", "Panel_leaderboard.prefab", "LocalLeaderboardService preload; backend thật là optional."),
    ("Quest", "Panel_quest.prefab", "Quest UI và QuestService local/config."),
    ("Daily Reward", "DailyRewardPopup.prefab", "Reward calendar; database được init từ module."),
    ("Feature Announcement", "UI Feature Announcement Popup.prefab", "Thông báo feature; không phải gameplay."),
    ("Lives/Add Lives", "UI Add Lives Panel.prefab", "Lives UI; LivesSystem là local/fallback."),
    ("Reward confirmation", "UI Rewards Confirmation Popup.prefab, ProgressPopUp.prefab", "Xác nhận reward/progress."),
    ("Network/no connection", "NoConnectionPopup.prefab / Internet Connection Popup.prefab", "Error/retry khi probe fail."),
    ("Dev Panel", "Dev Panel.prefab", "Debug tools; chỉ bật ở môi trường dev."),
], [1.35, 2.65, 2.55])
image_placeholder(6, "Menu popup hierarchy / flow", "Sơ đồ hoặc screenshot cho thấy bottom navigation mở Profile, Settings, Leaderboard, Quest.", "Có thể dùng ảnh Hierarchy + mũi tên flow; đánh dấu popup pause nếu có.", 1.6)
heading("3.4. Quy tắc popup và lỗi lifecycle đã sửa", 2)
para("Popup implement UIPage và có thể thêm IPopupWindow/IPausePopup. Khi mở popup pause, UIController theo dõi số popup và đặt Time.timeScale=0 nếu usePausePopups bật. Khi đóng, OnPageClosed khôi phục time scale khi count về 0.")
callout("LỖI ĐÃ XÁC ĐỊNH", "UIProfilePopup.OnDestroy() từng gọi BottomNavigationVisibilityEvents.RequestShow() lần thứ hai sau khi OnDisable đã phát event. Trong lúc scene/UI bị hủy, event thứ hai dẫn tới SetActive() trên No Ads button đang destroy. Quy tắc hiện tại: giữ xử lý ở OnDisable/CompleteHide; không phát event hiển thị bottom navigation trong OnDestroy.", RED, "B42318")
code("""// Mẫu đóng popup an toàn
public void Close()
{
    UIController.HidePage<UIProfilePopup>();
}

private void OnDisable()
{
    BottomNavigationVisibilityEvents.RequestShow();
}""")

# Part 4
page_break()
heading("4. Game scene: UI shell, session hook và Power Up bar", 1)
para("Game.unity hiện là gameplay shell. GameController khởi tạo UIController, mở UIGame và cung cấp hook cho complete/game-over/revive/replay/load menu. Scene không có board, tile, LevelController hoặc puzzle rules runtime; LevelData chỉ là contract dữ liệu nhẹ.")
image_placeholder(7, "Game scene HUD", "Cho thấy top HUD, timer/message area, safe area, Power Up bar và bottom navigation/banner.", "Chụp Game view khi Play Mode; nên có cả trạng thái Game, Complete và Game Over nếu có.", 1.85)
heading("4.1. GameController và trình tự vào game", 2)
bullet("Awake reset static flags, lấy UIController cùng GameObject, gọi Init()/InitPages() và AdsManager.EnableBanner().")
bullet("Start gọi UIController.ShowPage<UIGame>() rồi Overlay.Hide(0f).")
bullet("ActivateGame() là hook để game module báo đã sẵn sàng; SDK không tự tạo board.")
bullet("GameComplete() đánh dấu finish, gọi ActiveSession.Current.OnLevelCompleted(), QuestService.ReportProgress(CompleteLevels) và ShowCompleteUI().")
bullet("GameOver() mở UIGameOver; Revive() là hook provider/game module; Replay() và LoadMenu() điều hướng scene.")
code("""// Hook từ puzzle module (ý tưởng)
GameController.ActivateGame();
// Khi module tự xác định thắng:
GameController.GameComplete();
// Khi hết lượt/thua:
GameController.GameOver();""")
heading("4.2. UIGame và binding hiện có", 2)
para("UIGame kế thừa UIPage, có safeAreaRectTransform, TimerVisualiser gameplayTimer, MessageBox messageBox và PUUIController powerUpsUIController. Init() khởi tạo message box, Power Up UI và đăng ký safe area. PlayHideAnimation() ẩn timer rồi gọi UIController.OnPageClosed(this); PlayShowAnimation() gọi UIController.OnPageOpened(this).")
table(["Thành phần", "Prefab / field", "Cách dùng hiện tại"], [
    ("Safe area", "safeAreaRectTransform + NotchSaveArea", "Đặt HUD vào vùng an toàn; không hard-code notch offset."),
    ("Timer", "TimerVisualiser gameplayTimer", "Visual hook; timer thật phải do session/puzzle module cấp."),
    ("Message", "MessageBox messageBox", "Hiện trạng thái/hướng dẫn; Init trước khi gọi Show."),
    ("Power Up", "PUUIController + Power Up Panel.prefab", "UI-only preview; amount demo, không trừ inventory/không áp effect."),
    ("Complete", "UI Complete.prefab", "Hiện kết quả; nhận callback reward từ game/provider."),
    ("Game Over", "UI Game Over.prefab", "Hiện fail/revive/replay; logic revive ngoài SDK shell."),
], [1.25, 2.2, 3.1])
heading("4.3. Power Up bar: đủ UI, chưa có gameplay", 2)
para("PUUIController tạo bốn PUUIBehavior trong containerTransform từ itemPrefab, gán preview color và amount 1–4. Đây là cách giữ layout/active-deactive để team lồng gameplay sau này. Nó không biết PowerUpDefinition, inventory, cooldown, effect target hay board state.")
image_placeholder(8, "Power Up bar prefab / hierarchy", "Đối chiếu Power Up Panel, item prefab và các item active trong Game view.", "Chụp prefab mở rộng và Game view; ghi rõ đây là UI-only preview.", 1.65)
table(["Khi event", "UI shell nên làm", "Gameplay module phải làm"], [
    ("Session start", "Show bar, set amount/locked visual.", "Tạo session và cấp inventory thật."),
    ("Player tap Power Up", "Highlight/select, gọi callback.", "Validate cost, chọn target, áp effect."),
    ("Success/fail", "Refresh amount, hide selection.", "Commit state, analytics, save."),
    ("Scene unload", "Unsubscribe, clear references.", "Dispose board/session/effect runtime."),
], [1.45, 2.5, 2.6])
page_break()
heading("4.4. Luồng active/deactive đề xuất", 2)
number("GameController mở UIGame; UIGame.Init() bind UI một lần.")
number("Puzzle session gọi ActivateGame() sau khi board/input đã ready.")
number("PUUIController.Show() chỉ hiển thị các item mà config cho phép; không tự load LevelData.")
number("Khi pause/popup, UIController quản lý raycaster và Time.timeScale; Power Up bar không tự tắt nếu UX cần giữ HUD.")
number("Khi complete/game-over, hide input/Power Up callbacks trước, rồi mở UI Complete hoặc UI Game Over.")
number("Khi LoadMenu/reload scene, unsubscribe event và để page lifecycle kết thúc qua OnDisable/CompleteHide.")
callout("RANH GIỚI CONTRACT", "LevelData, Power Up và Game scene hiện là contract/UI shell. Không được thêm LevelController/Board/Tile giả chỉ để hết null; hãy để puzzle team cung cấp implementation qua interface ở phần 6.", AMBER, "B7791F")
heading("4.5. Checklist bind UI Game", 2)
table(["Hạng mục", "Đã có trong SDK", "Việc của game team"], [
    ("Canvas/safe area", "UIGame + NotchSaveArea", "Kéo đúng RectTransform, kiểm tra device notch."),
    ("Timer/message", "TimerVisualiser + MessageBox", "Cấp snapshot/session event; không tạo timer thứ hai trong UI."),
    ("Power Up", "PUUIController + item prefab", "Bind callback và inventory thật; giữ UI-only khi chưa có data."),
    ("Complete/fail", "UI Complete + UI Game Over", "Map Completed/Failed event từ session."),
    ("Scene exit", "UIController lifecycle", "Dispose board/session và unsubscribe trước khi destroy."),
], [1.35, 2.55, 2.65])
para("Nếu mọi ô trong checklist đều có owner rõ ràng, Game scene có thể chạy ổn định ngay cả khi puzzle engine được thay thế hoặc chưa được cài trong build SDK.")

# Part 5
page_break()
heading("5. Tổ chức module và cách dùng", 1)
para("Core module nằm dưới Assets/NebulaSoft Core/Modules; feature/project module nằm dưới Assets/Project Files/Game và Assets/Addon. Quy tắc phụ thuộc: UI gọi service abstraction; service không kéo ngược reference tới page cụ thể; provider optional đi qua adapter.")
heading("5.1. Cây thư mục khuyến nghị", 2)
code("""Assets/
├─ NebulaSoft Core/Modules/
│  ├─ Initializer, Save, Tween, Audio, Currency, Haptic, UI, Pool
│  ├─ Reward, Analytics, Skins, Monetization, Inspector, Defines
├─ Project Files/Game/
│  ├─ Scenes/Init.unity, Menu.unity, Game.unity
│  ├─ Scripts/Controllers, UI, Power Ups
│  ├─ Prefabs/UI, Lives System, Dev Panel
│  └─ Data/Project Init Settings.asset
└─ Addon/UI/
   ├─ Prefabs/Pages (Profile, Leaderboard, Quest, Daily Reward)
   └─ Prefabs/Shared (reward, connection, progress)""")
table(["Nhóm", "Module/folder hiện có", "Phụ thuộc chính"], [
    ("Core", "Initializer, Save, Tween, Audio, Currency, Haptic", "Unity runtime; config asset"),
    ("UI platform", "UI, Pool, Reward, Skins", "Canvas, prefab, service registry"),
    ("Game feature", "Daily Reward, Lives System, Quest, Leaderboard, Settings, Tutorial, Dev Panel", "Project Init Settings + UI prefab"),
    ("Integration", "Analytics, Monetization, Firebase, UI Store", "Provider adapter/SDK tùy chọn"),
    ("Puzzle gap", "Level System chỉ tối thiểu; không có Board/Tile/LevelController", "Game team phải cung cấp"),
], [1.3, 3.2, 2.05])
heading("5.2. Core modules", 2)
heading("Save", 3)
para("SaveInitModule gọi SaveController.Init(this). Dữ liệu bền vững nên đi qua save object/key của SaveController, không viết PlayerPrefs rải rác trong UI. autoSaveDelay hiện bằng 0 và cleanSaveStart tắt; WebGL prefix là build_20260223_173146.")
code("""// Ví dụ khái niệm: service tạo save object của feature
var profile = SaveController.GetSaveObject<PlayerProfileSave>("profile");
profile.LastSelectedLevel = levelId;
SaveController.Save();""")
para("Nếu API generic trong branch thay đổi, tra public signature của SaveController trước khi copy đoạn code; ví dụ trên mô tả boundary chứ không thay thế compile check.")
heading("Tween", 3)
para("TweenInitModule gọi Tween.Init() với 300 update, 30 fixed update, 0 late update. UI animation nên dùng Tween service thay vì coroutine tự quản lý khi cần pause/scene-safe.")
code("""Tween.Value(0f, 1f, 0.25f)
    .SetEase(Ease.OutCubic)
    .OnUpdate(value => fill.amount = value);""")
heading("Audio", 3)
para("AudioInitModule nạp AudioSettings, tạo pool 4 source, max distance 30 và spread 180. Menu/Game chỉ gọi abstraction để phát SFX/music; không tự new AudioSource mỗi click.")
heading("Currency và Reward", 3)
para("CurrencyInitModule nạp currenciesDatabase; Reward UI hiển thị confirmation/progress. UI có thể request reward nhưng commit currency phải qua service, rồi mới refresh label. Nếu provider economy chưa có, dùng local database/fallback.")
heading("Haptic", 3)
para("HapticInitModule chuẩn bị haptic service. Gọi haptic ở interaction semantic (tap success, error, reward), không gọi trong mỗi frame.")
heading("UI, Pool và Skins", 3)
para("UIController quản lý page/popup lifecycle; Pool dùng cho object lặp lại như FloatingCloud và cần prefab attached hợp lệ; Skins giữ mapping visual/config. Lỗi “no attached prefab at pool: FloatingCloud_Coins” nghĩa pool entry thiếu prefab, cần sửa asset pool chứ không tắt log.")
heading("Analytics", 3)
para("AnalyticsModules.Init() chạy trong Initializer.Init(), tách khỏi array Project Init Settings. Event nên phát sau khi state commit (level complete, reward claimed), không phát từ animation callback nếu có thể gọi lại.")
heading("5.3. Project feature và UI/local fallback", 2)
table(["Module", "Entry point / asset", "Ví dụ sử dụng"], [
    ("Game Settings", "GameInitModule → GameData.Init()", "Đọc game settings dùng chung; không chứa board runtime."),
    ("Lives System", "LivesSystemInitModule → LivesSystem.Init(LivesData)", "Add Lives popup, kiểm tra/consume lives qua service."),
    ("Quest", "QuestInitModule → QuestService.Init(QuestDatabase)", "ReportProgress(CompleteLevels) sau GameComplete."),
    ("Daily Reward", "DailyRewardInitModule → DailyRewardService.Init(database)", "Popup calendar đọc database; claim qua service."),
    ("Dev Panel", "DevPanelInitModule → DevPanelEnabler.LinkSettings(settings)", "Bật debug reset/telemetry trong dev build."),
    ("Leaderboard", "LocalLeaderboardService + Panel_leaderboard", "Preload local data; backend thật là optional."),
    ("Tutorial", "Tutorial Overlay / tutorial feature", "Overlay dẫn onboarding; không điều khiển puzzle rules."),
    ("Settings", "UI Menu Settings", "Đổi audio/haptic/screen; lưu qua service tương ứng."),
], [1.35, 2.9, 2.3])
heading("5.4. Optional provider và adapter", 2)
para("AdsManager, Firebase, IAP, remote config và cloud save không được trở thành compile-time dependency của Core. Tạo interface/adapter ở project integration, inject provider qua SDKInitializer hoặc một InitModule optional. Khi provider vắng, local fallback vẫn phải cho Menu/Game shell chạy.")
code("""public interface IAdsProvider
{
    void EnableBanner();
    void ShowRewarded(Action<bool> completed);
}

public sealed class AdsManagerAdapter : IAdsProvider
{
    // Bọc SDK thật; Core chỉ biết IAdsProvider.
}""")
image_placeholder(9, "Module configuration example in Inspector", "Có thể dùng hình này thêm lần nữa trong phần provider nếu cần minh họa adapter + config.", "Nếu tài liệu dài, thay bằng screenshot SDKInitializer và provider settings.", 1.45)

# Part 6
page_break()
heading("6. Contract-first: gắn puzzle runtime vào framework", 1)
para("Framework không áp đặt Picture Puzzle engine cũ. Game team cung cấp module puzzle qua contract; SDK chỉ cung cấp lifecycle, UI shell, navigation và service boundaries.")
heading("6.1. Các contract tối thiểu", 2)
table(["Contract", "Trách nhiệm puzzle module", "Framework hook"], [
    ("Level", "Load level id/config, metadata, difficulty.", "GameController/LevelData shell."),
    ("Board", "Tạo board, tile/node, layout và dispose.", "UIGame chỉ render vùng HUD; board nằm dưới layer riêng."),
    ("Input", "Tap/drag/gesture, lock khi popup/pause.", "UIController raycaster + session pause."),
    ("Session", "Start, pause, resume, dispose, elapsed time.", "TimerVisualiser nhận snapshot/session event."),
    ("Win/Fail", "Xác định điều kiện, score, stars, fail reason.", "GameComplete()/GameOver()."),
    ("Power Up", "Cost, inventory, target, cooldown, effect, result.", "PUUIController callback/visual only."),
    ("Telemetry", "Emit semantic events sau commit.", "AnalyticsModules/provider adapter."),
], [1.2, 3.35, 2.0])
heading("6.2. Mẫu interface đề xuất", 2)
code("""public interface IPuzzleSession
{
    string LevelId { get; }
    bool IsRunning { get; }
    float ElapsedSeconds { get; }
    void Start();
    void Pause();
    void Resume();
    void Dispose();
    event Action<PuzzleResult> Completed;
    event Action<PuzzleFailReason> Failed;
}

public interface IPuzzleBoard
{
    void Build(LevelData level);
    bool TryHandleInput(PuzzleInput input);
    void Dispose();
}""")
para("Đây là contract gợi ý, không phải class đã tồn tại trong SDK. Khi implement, đặt trong project puzzle module và map event về GameController ở một composition root của Game scene.")
heading("6.3. Composition root tại Game scene", 2)
number("Game scene có object PuzzleCompositionRoot do game team tạo; object này tham chiếu config/adapter, không sửa Core.")
number("Awake lấy service từ Initializer.InitSettings hoặc provider registry.")
number("Start load LevelData, build board, bind input và gọi session.Start().")
number("Session.Completed gọi GameController.GameComplete(); session.Failed gọi GameController.GameOver().")
number("OnDestroy/scene unload gọi session.Dispose(), board.Dispose() và unsubscribe mọi event.")
code("""public sealed class PuzzleCompositionRoot : MonoBehaviour
{
    [SerializeField] private PuzzleConfig config;
    private IPuzzleSession session;

    private void Start()
    {
        session = PuzzleFactory.Create(config);
        session.Completed += OnCompleted;
        session.Failed += OnFailed;
        session.Start();
        GameController.ActivateGame();
    }

    private void OnCompleted(PuzzleResult result) => GameController.GameComplete();
    private void OnFailed(PuzzleFailReason reason) => GameController.GameOver();

    private void OnDestroy()
    {
        if (session == null) return;
        session.Completed -= OnCompleted;
        session.Failed -= OnFailed;
        session.Dispose();
    }
}""")
callout("KHÔNG COPY ENGINE CŨ", "Không mang LevelController/Board/Tile cũ vào SDK chỉ để lấp khoảng trống. Nếu cần engine Picture Puzzle, đóng gói nó như một feature module implement các contract trên; framework vẫn chạy được khi module không cài.", AMBER, "B7791F")

# Part 7
page_break()
heading("7. Vận hành, kiểm tra và quy trình thay đổi", 1)
heading("7.1. Play Mode smoke check", 2)
table(["Bước", "Kỳ vọng"], [
    ("Play từ Init", "Loading hiển thị, không exception; Init → Menu."),
    ("Menu mở", "UIMainMenu hiện level map; bottom navigation và banner không spam lỗi."),
    ("Profile mở/đóng", "Bottom navigation ẩn/hiện đúng một lần; không SetActive object đang destroy."),
    ("Settings/Quest/Daily Reward", "Popup mở, đóng, time scale phục hồi; không duplicate page."),
    ("Nhấn Play", "SceneUtils xác nhận Game tồn tại; Game scene mở UIGame."),
    ("Game HUD", "Safe area, timer/message và Power Up bar hiển thị; Power Up chỉ là UI preview."),
    ("Complete/Game Over", "UI tương ứng mở; không giả định board/session có thật."),
    ("Đổi scene khi popup đang đóng", "Không có lifecycle exception trong Console."),
    ("Pool/Cloud", "Mọi pool entry có attached prefab; không còn lỗi FloatingCloud_Coins."),
], [2.2, 4.35])
heading("7.2. Kiểm tra Console theo nhóm lỗi", 2)
table(["Nhóm lỗi", "Cách khoanh vùng"], [
    ("GameObjects can not be made active when destroyed", "Tìm event gọi SetActive trong OnDestroy; kiểm tra OnDisable/CompleteHide có phát trùng."),
    ("Pool initialization failed", "Mở Pool asset, tìm key (ví dụ FloatingCloud_Coins), gắn prefab đúng field."),
    ("Missing reference UI", "Mở prefab/page, kiểm tra serialized field và scene override."),
    ("Provider unavailable", "Xác nhận fallback được bật; không xem warning provider là lỗi Core nếu feature optional."),
    ("Scene not found", "Build Settings có Init/Menu/Game và GameConsts.SCENE_GAME khớp."),
], [2.8, 3.75])
heading("7.3. Quy trình thêm prefab/UI", 2)
number("Đặt prefab dưới thư mục feature phù hợp; không đặt gameplay asset vào NebulaSoft Core.")
number("Kéo prefab vào UIController page hierarchy hoặc CachedPages theo pattern hiện có.")
number("Bổ sung CacheComponents()/Init() và safe area reference nếu page cần.")
number("Đăng ký popup interface nếu cần pause; kiểm tra OnPageOpened/OnPageClosed.")
number("Chạy smoke check mở/đóng, đổi scene và kiểm tra duplicate page.")
heading("7.4. Build scenes và release checklist", 2)
bullet("File → Build Settings có đúng thứ tự Init, Menu, Game; cả ba scene enabled.")
bullet("Init scene reference đúng ProjectInitSettings.asset và SDKInitializer.")
bullet("Mọi database/prefab trong modules array không null.")
bullet("Provider optional có adapter/fallback; build không phụ thuộc SDK không cài.")
bullet("No Ads, Leaderboard, Quest, Daily Reward và Power Up đều được kiểm tra theo đúng trạng thái UI-only/local fallback.")
bullet("Không tuyên bố framework có puzzle gameplay runtime nếu board/session/rules chưa được cung cấp.")

# Appendix
heading("Phụ lục A — Quick reference cho dev mới", 1)
table(["Muốn làm gì?", "Điểm bắt đầu"], [
    ("Thêm service chạy trước Menu", "InitModule + Project Init Settings asset"),
    ("Thêm popup", "UIPage prefab → UIController.InitPages → ShowPage<T>()"),
    ("Mở Game", "MenuController.OnPlayButtonClicked() → LoadGame()"),
    ("Hiển thị HUD", "UIGame + safeAreaRectTransform + child visualiser"),
    ("Thêm Power Up visual", "PUUIController/PUUIBehavior, không thêm data giả"),
    ("Gắn puzzle engine", "PuzzleCompositionRoot + IPuzzleSession/IPuzzleBoard"),
    ("Bật provider", "SDKInitializer/adapter optional, không sửa Core"),
    ("Xử lý lỗi pool", "Kiểm tra Pool entry và attached prefab"),
    ("Xử lý popup destroy", "OnDisable/CompleteHide một lần; tránh OnDestroy event"),
], [2.5, 4.05])
heading("Phụ lục B — Các đường dẫn cần nhớ", 1)
code("""Assets/Project Files/Game/Scenes/Init.unity
Assets/Project Files/Game/Scenes/Menu.unity
Assets/Project Files/Game/Scenes/Game.unity
Assets/Project Files/Data/Project Init Settings.asset
Assets/Project Files/Game/Scripts/Controllers/MenuController.cs
Assets/Project Files/Game/Scripts/Controllers/GameController.cs
Assets/Project Files/Game/Scripts/Power Ups/PUUIController.cs
Assets/Project Files/Game/Prefabs/UI/Canvas/UI Game.prefab
Assets/Project Files/Game/Prefabs/Power Ups/Power Up Panel.prefab
Assets/Addon/UI/Prefabs/Pages/UI Profile Popup.prefab""")
callout("BẢN GHI NHỚ CUỐI", "Khi bổ sung tính năng, xác định trước tính năng thuộc Core, project feature, provider optional hay puzzle module. Đặt code/config ở đúng tầng, expose contract nhỏ, và kiểm chứng cả đường vào scene lẫn đường thoát scene.", GREEN, "2E7D32")

doc.save(OUT)
print(OUT)
